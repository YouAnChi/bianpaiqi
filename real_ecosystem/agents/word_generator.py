import os
import json
import uvicorn
import logging
from pathlib import Path
from starlette.applications import Starlette
from starlette.responses import JSONResponse
from starlette.routing import Route
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API Key from environment
from dotenv import load_dotenv
load_dotenv()

if not os.getenv("GOOGLE_API_KEY"):
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

# Configure Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("WordGeneratorAgent")

class WordGeneratorAgent:
    def __init__(self):
        self.name = "Word Generator Agent"
        self.system_prompt = """
        你是一位专业的文档生成专家。
        你的任务是根据用户需求生成结构化的Word文档内容。
        你需要：
        1. 理解用户需求，确定文档的主题和结构
        2. 生成包含标题、段落、列表等结构化内容
        3. 内容应该专业、准确、格式清晰
        
        请以JSON格式返回文档结构，格式如下：
        {{
            "title": "文档标题",
            "sections": [
                {{
                    "heading": "章节标题",
                    "content": "章节内容...",
                    "subsections": [
                        {{
                            "heading": "子章节标题",
                            "content": "子章节内容..."
                        }}
                    ]
                }}
            ]
        }}
        
        请用中文生成内容。
        """
        self.llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.7)
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", self.system_prompt),
            ("user", "{input}")
        ])
        self.chain = self.prompt | self.llm | StrOutputParser()
        
        # Output directory
        self.output_dir = Path(os.getcwd()) / "output"
        self.output_dir.mkdir(exist_ok=True)

    def create_word_document(self, doc_structure: dict, filename: str) -> str:
        """根据结构化数据创建Word文档"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(doc_structure.get("title", "文档"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加章节
        sections = doc_structure.get("sections", [])
        for section in sections:
            # 一级标题
            doc.add_heading(section.get("heading", ""), level=1)
            
            # 章节内容
            content = section.get("content", "")
            if content:
                doc.add_paragraph(content)
            
            # 子章节
            subsections = section.get("subsections", [])
            for subsection in subsections:
                doc.add_heading(subsection.get("heading", ""), level=2)
                sub_content = subsection.get("content", "")
                if sub_content:
                    doc.add_paragraph(sub_content)
        
        # 保存文档
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info(f"Word document saved to: {filepath}")
        
        return str(filepath)

    async def handle_request(self, request):
        try:
            body = await request.json()
            logger.info(f"[{self.name}] Received request")
            
            # Extract user message
            user_msg = ""
            if 'params' in body and 'message' in body['params']:
                parts = body['params']['message'].get('parts', [])
                if parts:
                    user_msg = parts[0].get('text', '')
            
            if not user_msg:
                return JSONResponse({"error": "Empty message"}, status_code=400)

            # Try to parse user_msg as JSON (A2A protocol payload)
            task_desc = user_msg
            context = {}
            try:
                payload = json.loads(user_msg)
                if isinstance(payload, dict) and "task_description" in payload:
                    task_desc = payload.get("task_description", "")
                    context = payload.get("context", {})
                    
                    # Reconstruct a better prompt for the LLM
                    formatted_input = f"Task: {task_desc}\n\nContext:\n{json.dumps(context, indent=2, ensure_ascii=False)}"
                    user_msg = formatted_input
                    logger.info(f"[{self.name}] Parsed A2A JSON payload. Task: {task_desc[:50]}...")
            except json.JSONDecodeError:
                pass
            except Exception as e:
                logger.warning(f"[{self.name}] Error parsing payload: {e}")

            logger.info(f"[{self.name}] Processing: {user_msg[:50]}...")
            
            # Call LLM to generate document structure
            response_text = await self.chain.ainvoke({"input": user_msg})
            
            logger.info(f"[{self.name}] Response generated ({len(response_text)} chars)")
            
            # Parse JSON structure from response
            try:
                # Try to extract JSON from markdown code blocks
                if "```json" in response_text:
                    json_start = response_text.find("```json") + 7
                    json_end = response_text.find("```", json_start)
                    json_str = response_text[json_start:json_end].strip()
                elif "```" in response_text:
                    json_start = response_text.find("```") + 3
                    json_end = response_text.find("```", json_start)
                    json_str = response_text[json_start:json_end].strip()
                else:
                    json_str = response_text
                
                doc_structure = json.loads(json_str)
                
                # Generate filename
                safe_title = "".join(
                    [c for c in doc_structure.get("title", "document") if c.isalnum() or c in (' ', '-', '_')]
                ).strip().replace(' ', '_')[:50]
                filename = f"{safe_title}.docx"
                
                # Create Word document
                filepath = self.create_word_document(doc_structure, filename)
                
                result_text = f"Word文档已生成！\n\n文件路径: {filepath}\n\n文档标题: {doc_structure.get('title', '')}\n章节数: {len(doc_structure.get('sections', []))}"
                
            except json.JSONDecodeError as e:
                logger.warning(f"[{self.name}] Failed to parse JSON, using plain text format: {e}")
                # Fallback: create simple document from plain text
                doc = Document()
                doc.add_heading("生成的文档", level=0)
                doc.add_paragraph(response_text)
                
                filename = "generated_document.docx"
                filepath = self.output_dir / filename
                doc.save(str(filepath))
                
                result_text = f"Word文档已生成（简单格式）！\n\n文件路径: {filepath}\n\n内容:\n{response_text[:200]}..."

            # Construct A2A response
            response_data = {
                "result": {
                    "message": {
                        "role": "model",
                        "parts": [
                            {
                                "text": result_text
                            }
                        ]
                    }
                }
            }
            return JSONResponse(response_data)
            
        except Exception as e:
            logger.error(f"[{self.name}] Error: {e}", exc_info=True)
            return JSONResponse({"error": str(e)}, status_code=500)

agent = WordGeneratorAgent()
app = Starlette(debug=True, routes=[
    Route("/", agent.handle_request, methods=["POST"]),
])

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=10008)
